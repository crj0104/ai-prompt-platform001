from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

OUT = r"F:\数据库原理课设\deliverables\AI提示词模板平台课程设计报告.docx"


def set_run_font(run, size=None, bold=None, color=None, font="宋体"):
    run.font.name = font
    run._element.rPr.rFonts.set(qn("w:eastAsia"), font)
    run._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    run._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    if size is not None:
        run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold
    if color is not None:
        run.font.color.rgb = RGBColor.from_string(color)


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_text(cell, text, bold=False, size=10.5, color=None, align=WD_ALIGN_PARAGRAPH.LEFT):
    cell.text = ""
    p = cell.paragraphs[0]
    p.alignment = align
    p.paragraph_format.space_after = Pt(0)
    run = p.add_run(text)
    set_run_font(run, size=size, bold=bold, color=color)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def set_table_borders(table, color="D9DEE8"):
    tbl = table._tbl
    tbl_pr = tbl.tblPr
    borders = tbl_pr.first_child_found_in("w:tblBorders")
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        tag = f"w:{edge}"
        element = borders.find(qn(tag))
        if element is None:
            element = OxmlElement(tag)
            borders.append(element)
        element.set(qn("w:val"), "single")
        element.set(qn("w:sz"), "6")
        element.set(qn("w:space"), "0")
        element.set(qn("w:color"), color)


def add_heading(doc, text, level=1):
    p = doc.add_heading(level=level)
    p.paragraph_format.space_before = Pt(16 if level == 1 else 10)
    p.paragraph_format.space_after = Pt(7 if level == 1 else 5)
    run = p.add_run(text)
    set_run_font(run, size={1: 16, 2: 13, 3: 12}.get(level, 11), bold=True, color="2E74B5" if level < 3 else "1F4D78")
    return p


def add_para(doc, text="", first_line=True):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.line_spacing = 1.1
    if first_line and text:
        p.paragraph_format.first_line_indent = Inches(0.28)
    run = p.add_run(text)
    set_run_font(run, size=11)
    return p


def add_bullet(doc, text):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run(text)
    set_run_font(run, size=11)
    return p


def add_number(doc, text):
    p = doc.add_paragraph(style="List Number")
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run(text)
    set_run_font(run, size=11)
    return p


def add_caption(doc, text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(3)
    p.paragraph_format.space_after = Pt(8)
    run = p.add_run(text)
    set_run_font(run, size=9.5, color="555555")


def add_table(doc, headers, rows, widths=None):
    table = doc.add_table(rows=1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    set_table_borders(table)
    for i, h in enumerate(headers):
        set_cell_text(table.rows[0].cells[i], h, bold=True, color="1F4D78", align=WD_ALIGN_PARAGRAPH.CENTER)
        set_cell_shading(table.rows[0].cells[i], "F2F4F7")
    for row in rows:
        cells = table.add_row().cells
        for i, value in enumerate(row):
            align = WD_ALIGN_PARAGRAPH.CENTER if i == 0 and len(headers) > 2 else WD_ALIGN_PARAGRAPH.LEFT
            set_cell_text(cells[i], str(value), align=align)
    if widths:
        for row in table.rows:
            for cell, width in zip(row.cells, widths):
                cell.width = Inches(width)
    return table


def add_code(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.2)
    p.paragraph_format.right_indent = Inches(0.2)
    p.paragraph_format.space_after = Pt(6)
    run = p.add_run(text)
    run.font.name = "Consolas"
    run._element.rPr.rFonts.set(qn("w:ascii"), "Consolas")
    run._element.rPr.rFonts.set(qn("w:hAnsi"), "Consolas")
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    run.font.size = Pt(9.5)
    return p


def add_diagram_table(doc, headers, rows, caption):
    table = add_table(doc, headers, rows)
    add_caption(doc, caption)
    return table


doc = Document()
section = doc.sections[0]
section.top_margin = Inches(1)
section.bottom_margin = Inches(1)
section.left_margin = Inches(1)
section.right_margin = Inches(1)

styles = doc.styles
normal = styles["Normal"]
normal.font.name = "宋体"
normal._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
normal.font.size = Pt(11)

header = section.header.paragraphs[0]
header.alignment = WD_ALIGN_PARAGRAPH.CENTER
set_run_font(header.add_run("AI提示词模板平台课程设计报告"), size=9.5, color="666666")
footer = section.footer.paragraphs[0]
footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
set_run_font(footer.add_run("数据库原理课程设计"), size=9.5, color="666666")

# Cover
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.space_before = Pt(88)
r = p.add_run("AI提示词模板平台\n数据库应用系统课程设计报告")
set_run_font(r, size=24, bold=True, color="0B2545", font="黑体")

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.space_before = Pt(28)
r = p.add_run("课题类型：数据库应用系统开发\n技术路线：Spring Boot + MyBatis-Plus + H2/MySQL + 原生前端\n提交材料：课程设计报告、汇报PPT、源代码")
set_run_font(r, size=13, color="333333")

table = doc.add_table(rows=4, cols=2)
table.alignment = WD_TABLE_ALIGNMENT.CENTER
set_table_borders(table, "B8BCC4")
meta = [("小组成员", "成员A、成员B（提交前替换为真实姓名）"), ("指导教师", "________________"), ("班级", "________________"), ("完成日期", "2026年7月")]
for row, item in zip(table.rows, meta):
    set_cell_text(row.cells[0], item[0], bold=True, align=WD_ALIGN_PARAGRAPH.CENTER)
    set_cell_text(row.cells[1], item[1])
doc.add_page_break()

add_heading(doc, "目录", 1)
toc_items = [
    "第一章 课题描述",
    "第二章 系统分析",
    "第三章 数据库设计",
    "第四章 系统设计与实现",
    "第五章 总结",
    "参考文献",
]
for item in toc_items:
    add_para(doc, item, first_line=False)
doc.add_page_break()

add_heading(doc, "第一章 课题描述", 1)
add_heading(doc, "1.1 课题概述", 2)
add_para(doc, "本课题围绕“AI提示词模板平台”构建一个数据库应用系统。系统面向普通用户、创作者和平台管理员三类角色，提供提示词模板检索、模板详情查看、购买与领取、使用记录、评分评价、收藏管理、创作者看板、平台统计等功能。项目采用前后端分离方式组织，后端位于 backend 目录，前端位于 frontend 目录，数据库脚本位于 backend/src/main/resources/sql。")
add_para(doc, "AI提示词模板平台的业务特点在于模板内容具有版本演进、模板需要分类检索、付费模板需要交易与权限控制、使用和评价需要保留行为记录。因此系统的数据库设计不仅要满足基础增删改查，还要体现版本管理、交易一致性、统计分析、查询优化和并发控制等数据库课程设计要求。")

add_heading(doc, "1.2 设计要求", 2)
for text in [
    "完成数据库应用系统的需求分析，明确系统参与者、业务用例和功能边界。",
    "完成数据库概念模型、逻辑模型和物理模型设计，并用 E-R 图或结构图说明实体关系。",
    "完成系统功能模块结构设计和核心代码实现，能够进行系统演示。",
    "完成并发控制实验与分析、查询优化实验与分析，体现数据库设计与实现能力。",
    "完成课程设计报告、汇报 PPT 和源代码整理，最终按小组人员姓名压缩提交。",
]:
    add_bullet(doc, text)

add_heading(doc, "1.3 设计目标", 2)
for text in [
    "实现一个可运行、可演示、结构清晰的 AI 提示词模板交易与管理平台。",
    "设计合理的数据库表结构，覆盖用户、模板、版本、标签、收藏、订单、使用日志、评价、限免活动等核心数据。",
    "通过统一服务编排层降低控制器与领域服务耦合，提高项目结构可读性。",
    "通过索引、分页查询、聚合统计和事务控制提升系统性能与数据一致性。",
    "通过前后端分离方式展示现代 Web 应用开发流程。",
]:
    add_bullet(doc, text)

add_heading(doc, "1.4 小组人员分工", 2)
add_table(
    doc,
    ["成员", "主要职责", "具体工作"],
    [
        ["成员A", "数据库与后端开发", "需求分析、E-R模型设计、建表脚本、实体类与 Mapper、交易与使用流程、事务控制、后端测试"],
        ["成员B", "前端与文档测试", "前端页面设计、接口联调、系统演示准备、查询优化实验记录、课程设计报告和PPT整理"],
    ],
    [1.0, 1.7, 3.8],
)
add_caption(doc, "表 1-1 小组人员分工")

add_heading(doc, "第二章 系统分析", 1)
add_heading(doc, "2.1 需求分析", 2)
add_para(doc, "系统需要支持用户从模板市场中检索合适的提示词模板，并根据模板价格类型进行免费使用或付费购买。创作者需要查看已发布模板的使用量、版本变更、收益趋势等信息。平台管理员需要查看平台整体运营概况，包括用户数、模板数、交易额、热门模板和限免活动。")
add_table(
    doc,
    ["角色", "核心需求", "对应功能"],
    [
        ["普通用户", "查找并使用高质量提示词模板", "注册登录、模板搜索、详情查看、收藏、购买、使用、评价、个人中心"],
        ["创作者", "管理模板并了解作品表现", "模板列表、版本信息、使用趋势、收入趋势、创作者等级展示"],
        ["平台管理员", "掌握平台经营与活动状态", "平台统计、热门排行、限免活动、交易与使用数据汇总"],
    ],
    [1.1, 2.3, 3.1],
)
add_caption(doc, "表 2-1 用户角色与需求映射")

add_heading(doc, "2.2 用例图", 2)
use_rows = [
    ["普通用户", "注册/登录、搜索模板、查看详情、收藏模板、购买模板、使用模板、提交评价、查看个人中心"],
    ["创作者", "查看已发布模板、查看使用趋势、查看月度收入、查看版本历史"],
    ["管理员", "查看平台总览、查看热门模板、查看限免活动、查看统计指标"],
    ["系统", "校验付费权限、记录使用日志、生成订单、维护冗余统计字段、执行分页与聚合查询"],
]
add_diagram_table(doc, ["参与者", "用例集合"], use_rows, "图 2-1 系统用例图（文本化表示）")

add_heading(doc, "2.3 功能模块划分", 2)
module_rows = [
    ["用户中心模块", "注册、登录、个人资料、订单、余额流水、使用记录"],
    ["模板领域模块", "模板搜索、模板详情、标签筛选、推荐模板、版本展示"],
    ["交易使用模块", "收藏/取消收藏、购买、使用模板、评价、余额流水"],
    ["统计分析模块", "使用趋势、热门排行、创作者看板、平台总览"],
    ["前端展示模块", "模板市场、详情面板、个人中心、创作者看板、平台看板"],
]
add_diagram_table(doc, ["功能模块", "包含功能"], module_rows, "图 2-2 功能模块结构图")

add_heading(doc, "2.4 非功能需求", 2)
for text in [
    "一致性：购买、扣款、订单、余额流水应处于同一事务边界内，避免部分成功。",
    "可扩展性：通过 PortalService 统一门面和领域服务拆分，便于后续接入真实登录态、模板编辑、后台管理。",
    "查询效率：模板搜索使用分页查询，热门排行和趋势统计下推到 SQL 聚合，并配合索引设计。",
    "可演示性：默认使用 H2 内存数据库，项目启动后无需安装 MySQL 即可完成演示。",
]:
    add_bullet(doc, text)

add_heading(doc, "第三章 数据库设计", 1)
add_heading(doc, "3.1 概念模型设计", 2)
add_para(doc, "根据业务需求，系统抽象出用户、模板、模板版本、标签、订单、余额流水、使用日志、评价、限免活动、活动模板等实体。用户与模板之间存在创作、收藏、购买、使用、评价等多种关系；模板与模板版本为一对多关系；模板与标签为多对多关系；订单、使用日志和评价共同构成“购买-使用-反馈”的业务闭环。")
er_rows = [
    ["SysUser 用户", "1:N 创作", "PromptTemplate 模板"],
    ["PromptTemplate 模板", "1:N 拥有", "PromptTemplateVersion 模板版本"],
    ["PromptTemplate 模板", "N:M 归类", "Tag 标签（通过 template_tag_rel）"],
    ["SysUser 用户", "N:M 收藏", "PromptTemplate 模板（通过 template_favorite）"],
    ["SysUser 用户", "1:N 下单", "TemplateOrder 订单"],
    ["TemplateOrder 订单", "0/1:N 关联", "TemplateUseLog 使用日志"],
    ["TemplateUseLog 使用日志", "1:0/1 产生", "TemplateReview 评价"],
    ["FreeCampaign 活动", "1:N 包含", "CampaignTemplate 活动模板"],
]
add_diagram_table(doc, ["实体A", "关系", "实体B"], er_rows, "图 3-1 E-R 概念模型关系图")

add_heading(doc, "3.2 逻辑模型设计", 2)
logic_rows = [
    ["sys_user", "user_id", "用户账号、联系方式、创作者等级、余额"],
    ["prompt_template", "template_id", "模板主表，保存标题、场景、价格、当前版本、统计字段"],
    ["prompt_template_version", "version_id", "模板版本表，保存提示词内容、版本号、变更说明"],
    ["tag", "tag_id", "标签表，支持父子标签和标签路径"],
    ["template_tag_rel", "id", "模板与标签多对多关系"],
    ["template_favorite", "favorite_id", "用户收藏关系"],
    ["template_order", "order_id", "模板购买订单"],
    ["user_balance_log", "log_id", "余额变动流水与审计记录"],
    ["template_use_log", "use_log_id", "模板使用记录"],
    ["template_review", "review_id", "模板评分评价"],
    ["free_campaign", "campaign_id", "限时免费活动"],
    ["campaign_template", "campaign_template_id", "活动与模板关系及库存"],
    ["campaign_claim", "claim_id", "活动领取记录，防重复领取"],
]
add_table(doc, ["表名", "主键", "说明"], logic_rows, [1.85, 1.5, 3.15])
add_caption(doc, "表 3-1 逻辑模型表设计")

add_heading(doc, "3.3 关系约束设计", 2)
for text in [
    "prompt_template.creator_user_id 外键引用 sys_user.user_id，表示模板创作者。",
    "prompt_template.current_version_id 外键引用 prompt_template_version.version_id，表示当前生效版本。",
    "prompt_template_version 通过 (template_id, version_no) 唯一约束避免同一模板版本号重复。",
    "template_tag_rel 通过 (template_id, tag_id) 唯一约束避免重复打标。",
    "template_favorite 通过 (user_id, template_id) 唯一约束避免重复收藏。",
    "template_review.score 设置 1 到 5 的检查约束，保证评分范围合法。",
    "campaign_claim 通过 (campaign_template_id, user_id) 唯一约束防止同一用户重复领取同一活动模板。",
]:
    add_bullet(doc, text)

add_heading(doc, "3.4 物理模型设计", 2)
physical_rows = [
    ["idx_template_filter", "prompt_template(shelf_status, avg_score, price)", "支持模板筛选和评分/价格过滤"],
    ["idx_template_hot_sort", "prompt_template(shelf_status, use_count, favorite_count, avg_score)", "支持热门模板排序"],
    ["idx_template_creator", "prompt_template(creator_user_id, created_at)", "支持创作者模板列表"],
    ["idx_rel_tag_template", "template_tag_rel(tag_id, template_id)", "支持按标签查模板"],
    ["idx_order_query", "template_order(user_id, pay_status, created_at)", "支持用户订单查询"],
    ["idx_use_query", "template_use_log(template_id, used_at)", "支持模板使用趋势统计"],
    ["idx_use_user_time", "template_use_log(user_id, used_at)", "支持用户使用记录查询"],
    ["idx_review_query", "template_review(template_id, score, created_at)", "支持评价统计与展示"],
    ["idx_campaign_time", "free_campaign(start_time, end_time, campaign_status)", "支持活动有效期查询"],
]
add_table(doc, ["索引名", "字段", "作用"], physical_rows, [1.55, 2.8, 2.15])
add_caption(doc, "表 3-2 物理模型索引设计")
add_para(doc, "系统演示环境使用 H2 内存数据库，并开启 MySQL 兼容模式；真实部署时可执行 schema-mysql.sql 创建 MySQL 8.0 数据库。金额字段采用 DECIMAL(10,2)，时间字段采用 DATETIME/TIMESTAMP，提示词内容使用 TEXT/CLOB，以满足较长模板内容存储。")

add_heading(doc, "第四章 系统设计与实现", 1)
add_heading(doc, "4.1 总体架构设计", 2)
add_para(doc, "系统采用前后端分离架构。前端使用原生 HTML、CSS、JavaScript 实现，负责页面交互与 API 调用；后端使用 Spring Boot 提供 REST API；持久层使用 MyBatis-Plus 和 XML Mapper；数据库层使用 H2/MySQL 保存业务数据。")
arch_rows = [
    ["前端层", "frontend/index.html、styles.css、app.js", "模板市场、详情、个人中心、看板页面"],
    ["接口层", "ApiController", "统一暴露 /api 下的 REST 接口"],
    ["编排层", "PortalService / PortalApplicationService", "聚合领域服务，对 Controller 提供统一入口"],
    ["领域服务层", "TemplateDomainService、TradeService、UserCenterService 等", "实现模板、交易、用户、统计等业务规则"],
    ["持久层", "Mapper + XML SQL", "分页查询、聚合统计、实体持久化"],
    ["数据库层", "H2/MySQL", "业务数据存储、约束、索引和事务"],
]
add_diagram_table(doc, ["层次", "实现", "职责"], arch_rows, "图 4-1 系统总体架构图")

add_heading(doc, "4.2 后端统一编排设计", 2)
add_para(doc, "后端控制器只依赖 PortalService 接口，具体实现类为 PortalApplicationService。该服务不直接承载所有业务细节，而是聚合模板领域服务、交易服务、用户中心服务、统计服务等领域服务，形成类似应用服务层的统一编排入口。")
add_code(doc, "ApiController -> PortalService -> PortalApplicationService -> Domain Services -> Mapper -> Database")
add_para(doc, "该设计的优点是接口层稳定、业务边界清晰。后续如果接入真实登录态，只需在 PortalApplicationService 中替换当前演示用的 DEFAULT_USER_ID 获取方式，不必修改各个 Controller。")

add_heading(doc, "4.3 主要功能模块实现", 2)
add_table(
    doc,
    ["模块", "关键类/文件", "实现说明"],
    [
        ["模板市场", "TemplateDomainServiceImpl、PromptTemplateMapper.xml", "支持关键字、标签、评分、价格、排序和分页检索"],
        ["模板详情", "TemplateDomainServiceImpl", "展示当前版本内容、版本历史、标签、评价和统计数据"],
        ["购买与使用", "TradeServiceImpl", "校验模板状态、重复购买、付费权限，记录订单和使用日志"],
        ["收藏评价", "TradeServiceImpl", "收藏/取消收藏维护关系表与冗余计数，评价前校验使用记录"],
        ["用户中心", "UserCenterServiceImpl", "注册、登录校验、个人资料、订单、使用记录、余额流水"],
        ["统计看板", "TemplateAnalyticsServiceImpl、DashboardServiceImpl", "使用趋势、热门排行、创作者和平台统计"],
        ["前端页面", "frontend/app.js", "通过 fetch 调用后端 API，完成搜索、购买、使用、评价等交互"],
    ],
    [1.2, 2.0, 3.3],
)
add_caption(doc, "表 4-1 功能模块实现说明")

add_heading(doc, "4.4 关键业务流程", 2)
purchase_rows = [
    ["1", "读取模板与用户", "校验模板存在、用户存在、模板已上架"],
    ["2", "判断重复购买", "若已有 SUCCESS 订单，则直接返回已购买结果，避免重复扣款"],
    ["3", "校验余额", "付费模板要求用户余额不低于模板价格"],
    ["4", "生成订单", "写入 template_order，记录 pay_status 和 order_status"],
    ["5", "扣减余额并记流水", "更新 sys_user.balance，写入 user_balance_log"],
    ["6", "返回结果", "返回订单视图和最新余额"],
]
add_diagram_table(doc, ["步骤", "处理", "说明"], purchase_rows, "图 4-2 模板购买流程图")

use_rows = [
    ["1", "用户点击使用模板", "前端提交 inputSummary"],
    ["2", "后端校验模板状态", "模板必须存在并且处于 ON_SHELF"],
    ["3", "校验付费权限", "免费模板可直接使用；付费模板必须存在成功订单"],
    ["4", "读取当前版本", "优先使用 prompt_template.current_version_id"],
    ["5", "写入使用日志", "记录用户、模板、版本、订单、输入摘要和来源"],
    ["6", "更新使用计数", "use_count 原子递增"],
]
add_diagram_table(doc, ["步骤", "处理", "说明"], use_rows, "图 4-3 模板使用流程图")

add_heading(doc, "4.5 前端界面实现", 2)
add_para(doc, "前端放置于 frontend 目录，主要由 index.html、styles.css 和 app.js 组成。页面左侧提供视图导航和后端 API 地址配置，右侧主区域按视图展示模板市场、个人中心、创作者看板和平台总览。")
for text in [
    "模板市场：支持关键字、标签、排序、最低评分筛选，列表中显示标题、场景、标签、价格、评分和使用量。",
    "模板详情：展示提示词内容、作者、评分、收藏数、评价，并提供购买、收藏、使用和评价按钮。",
    "个人中心：展示用户余额、订单、使用记录和余额流水。",
    "统计看板：展示创作者统计和平台统计，体现后端聚合查询结果。",
]:
    add_bullet(doc, text)

add_heading(doc, "4.6 并发控制实验与分析", 2)
add_para(doc, "交易类系统容易出现重复扣款、重复购买、计数覆盖等并发问题。本项目在优化过程中对购买、收藏和使用流程进行了并发风险分析。")
add_table(
    doc,
    ["实验场景", "潜在问题", "处理方式", "分析结论"],
    [
        ["同一用户重复购买同一模板", "短时间多次点击可能重复生成订单和重复扣款", "购买前查询 SUCCESS 订单，已购买则直接返回已有订单", "提高了接口幂等性，适合演示与基础并发保护"],
        ["多人同时收藏同一模板", "读取 favorite_count 后再写回可能丢失更新", "使用 SQL 表达式 favorite_count = favorite_count + 1/-1", "避免简单读改写导致的计数覆盖"],
        ["多人同时使用模板", "use_count 可能统计不准确", "使用 use_count = use_count + 1 原子递增", "计数更新更符合高频使用场景"],
        ["余额扣减", "真实高并发下可能出现超扣风险", "当前使用事务保证流程原子性；后续可改为带余额条件的 UPDATE", "课程设计阶段可说明事务边界，真实系统应进一步加强"],
    ],
    [1.45, 1.65, 1.7, 1.7],
)
add_caption(doc, "表 4-2 并发控制实验与分析")
add_para(doc, "实验结论：在 Spring 事务控制下，订单、余额、流水等操作可作为一个整体提交或回滚；对收藏数、使用数等高频计数字段采用 SQL 原子递增可以降低丢失更新概率。后续若部署到真实 MySQL，应进一步通过行级锁、乐观锁版本号或条件更新语句增强余额扣减的并发安全。")

add_heading(doc, "4.7 查询优化实验与分析", 2)
add_para(doc, "系统对模板搜索、热门排行、使用趋势等高频查询进行了优化。优化思路主要包括：将筛选条件下推到 SQL、使用分页插件、增加组合索引、使用聚合查询而不是内存遍历。")
add_table(
    doc,
    ["查询场景", "优化前风险", "当前实现", "涉及索引/SQL"],
    [
        ["模板搜索", "先查全量再内存过滤会造成分页不准和性能下降", "PromptTemplateMapper.xml 中按条件 SQL 过滤并分页", "idx_template_filter、idx_rel_tag_template"],
        ["热门排行", "多维热度计算如果全部在 Java 内存处理会浪费资源", "SQL 统计近期使用量并计算 hotScore", "idx_template_hot_sort、idx_use_query"],
        ["使用趋势", "逐条读取日志后按日期分组效率低", "TemplateUseLogMapper.xml 按日期 GROUP BY", "idx_use_query"],
        ["个人订单/记录", "按用户查询如果缺少索引会全表扫描", "按 user_id + 时间排序查询", "idx_order_query、idx_use_user_time"],
    ],
    [1.3, 1.9, 1.8, 1.5],
)
add_caption(doc, "表 4-3 查询优化实验与分析")
add_para(doc, "测试过程中通过 mvn test 验证模板搜索、详情查询、登录校验、购买幂等、付费模板使用权限、取消收藏等接口。当前测试共 7 个用例通过，说明核心查询和业务流程可以在 H2 演示数据库下稳定运行。")

add_heading(doc, "4.8 开发环境与关键技术", 2)
add_table(
    doc,
    ["类别", "技术/工具", "作用"],
    [
        ["后端框架", "Spring Boot 3.5.3", "快速构建 REST API 服务"],
        ["持久层", "MyBatis-Plus", "实体映射、分页插件、Mapper 简化"],
        ["数据库", "H2 / MySQL 8.0", "演示数据库与真实部署数据库"],
        ["前端", "HTML + CSS + JavaScript", "独立前端页面与 API 调用"],
        ["构建工具", "Maven", "依赖管理、测试和启动"],
        ["测试", "Spring Boot Test + MockMvc", "接口级自动化验证"],
    ],
    [1.4, 2.1, 3.0],
)
add_caption(doc, "表 4-4 开发环境与关键技术")

add_heading(doc, "第五章 总结", 1)
add_heading(doc, "5.1 实现过程碰到的问题及解决", 2)
add_table(
    doc,
    ["问题", "表现", "解决方法"],
    [
        ["前后端耦合", "原项目中 Thymeleaf 页面与后端在同一工程中", "拆分为 backend 和 frontend 两个目录，后端只保留 REST API"],
        ["业务流程分散", "Controller 如果直接调用多个领域服务会导致代码耦合", "保留 PortalService 统一门面，由 PortalApplicationService 编排流程"],
        ["付费模板权限缺失", "未购买用户可能直接使用付费模板", "TradeServiceImpl 中使用前校验成功订单"],
        ["重复购买风险", "同一用户多次购买同一模板可能重复扣款", "购买前查询已有成功订单，接口具备幂等性"],
        ["当前版本不一致", "按创建时间取版本可能与 current_version_id 不一致", "统一优先读取 current_version_id 指向版本"],
        ["查询性能", "模板搜索、趋势统计若全部内存处理会影响性能", "使用 SQL 条件过滤、聚合查询和组合索引"],
    ],
    [1.45, 2.0, 3.05],
)
add_caption(doc, "表 5-1 编程与调试问题总结")

add_heading(doc, "5.2 系统完成情况", 2)
for text in [
    "完成前后端分离目录结构，后端 API 与前端静态页面可独立维护。",
    "完成核心数据库表设计、外键约束、唯一约束和索引设计。",
    "完成模板搜索、详情、收藏、购买、使用、评价、个人中心和统计看板功能。",
    "完成统一编排层和领域服务拆分，代码结构较清晰。",
    "完成后端接口测试，核心测试用例通过。",
]:
    add_bullet(doc, text)

add_heading(doc, "5.3 展望与可改进之处", 2)
for text in [
    "接入 Spring Security 或 JWT，实现真实用户登录态，替换当前演示用 DEFAULT_USER_ID。",
    "增加模板创建、编辑、审核、上下架和版本回滚页面，使创作者功能更完整。",
    "余额扣减进一步使用条件更新或乐观锁，增强真实高并发支付场景下的安全性。",
    "完善 campaign_claim 的实体、Mapper 和活动领取流程，使限免活动形成完整闭环。",
    "将提示词内容接入向量检索或语义搜索，提高模板推荐质量。",
    "补充更多自动化测试和压力测试，验证并发购买、收藏和使用场景。",
]:
    add_bullet(doc, text)

add_heading(doc, "5.4 提交说明", 2)
add_para(doc, "课程设计检查时，小组需提交一份课程设计报告、一份汇报 PPT，并进行系统演示。检查完毕后，将课程设计报告、汇报 PPT 和源代码一起压缩，压缩包文件名建议为“小组人员姓名.zip”，例如“张三_李四.zip”。")

add_heading(doc, "参考文献", 1)
refs = [
    "[1] Abraham Silberschatz, Henry F. Korth, S. Sudarshan. Database System Concepts. McGraw-Hill Education.",
    "[2] 王珊, 萨师煊. 数据库系统概论. 高等教育出版社.",
    "[3] MySQL 8.0 Reference Manual. Oracle Corporation.",
    "[4] Spring Boot Reference Documentation. VMware, Inc.",
    "[5] MyBatis-Plus 官方文档. https://baomidou.com/",
    "[6] MDN Web Docs. HTML、CSS、Fetch API Documentation.",
]
for ref in refs:
    add_para(doc, ref, first_line=False)

doc.save(OUT)
print(OUT)
